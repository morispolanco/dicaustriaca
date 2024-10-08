import streamlit as st
import requests
import json
from docx import Document
from io import BytesIO

# Set page configuration
st.set_page_config(page_title="Diccionario Económico Austríaco", page_icon="📚", layout="wide")

# Function to create the information column
def crear_columna_info():
    st.markdown("""
    ## Sobre esta aplicación

    Esta aplicación es un Diccionario Económico basado en la perspectiva de la Escuela Austríaca de Economía. Permite a los usuarios obtener definiciones de términos económicos según esta interpretación.

    ### Cómo usar la aplicación:

    1. Elija un término económico de la lista predefinida o ingrese su propio término.
    2. Haga clic en "Generar entrada de diccionario" para obtener la definición desde la perspectiva de la escuela austríaca.
    3. Lea la definición y las fuentes proporcionadas.
    4. Si lo desea, descargue un documento DOCX con toda la información.

    ### Autor y actualización:
    **Moris Polanco**, 28 ag 2024

    ### Cómo citar esta aplicación (formato APA):
    Polanco, M. (2024). *Diccionario Económico Austríaco* [Aplicación web]. https://escuelaaustriaca.streamlit.app

    ---
    **Nota:** Esta aplicación utiliza inteligencia artificial para generar definiciones basadas en la visión de la escuela austríaca. Verifique la información con fuentes adicionales para un análisis más profundo.
    """)

# Titles and Main Column
st.title("Diccionario Económico Austríaco")

col1, col2 = st.columns([1, 2])

with col1:
    crear_columna_info()

with col2:
    TOGETHER_API_KEY = st.secrets["TOGETHER_API_KEY"]
    SERPLY_API_KEY = st.secrets["SERPLY_API_KEY"]

    # 101 economic terms related to the Austrian school of economics
    terminos_economicos = sorted([
        "Acción Humana", "Ahorro", "Aranceles", "Armonía Económica", "Avería", "Banco Central",
        "Bienes de Capital", "Bienes Intermedios", "Bienes de Consumo", "Capitalismo", "Competencia",
        "Competencia Monopolística", "Competencia Perfecta", "Conocimiento", "Costo de Oportunidad",
        "Crédito", "Crecimiento Económico", "Ciclo Económico", "Deflación", "Demanda", "División del Trabajo",
        "Doble Coincidencia de Deseos", "Eficiencia", "Elasticidad", "Emprendimiento", "Equilibrio Económico",
        "Especialización", "Espontaneidad", "Esperanza de Vida", "Estado de Derecho", "Externalidades",
        "Factor de Producción", "Federalismo", "Fiduciario", "Función Empresarial", "Futuro", "Gasto Público",
        "Inflación", "Instituciones", "Interés", "Inversión", "Intervencionismo", "Libre Mercado",
        "Mecanismo de Precios", "Mercado", "Microeconomía", "Modelo de Competencia", "Moneda",
        "Monopolio", "Oferta", "Orden Espontáneo", "Paradigma", "Pareto", "Plusvalía", "Poder Adquisitivo",
        "Política Económica", "Ponderación", "Precio", "Preferencia de Tiempo", "Preferencias", "Producción",
        "Productividad", "Propiedad Privada", "Proteccionismo", "Racionalidad", "Recurso Económico",
        "Redistribución de la Riqueza", "Regulación", "Renta", "Riesgo", "Sector Público", "Sector Privado",
        "Seguridad Jurídica", "Servicio", "Sistema Económico", "Soberanía del Consumidor", "Sociedad Abierta",
        "Subsidio", "Sujeto Económico", "Tasa de Interés", "Teoría del Ciclo Económico", "Trabajo", "Valor",
        "Valor de Uso", "Valor del Cambio", "Ventaja Competitiva", "Ventaja Comparativa", "Verosimilitud",
        "Voluntad Individual", "Bienes Públicos", "Economía de Escala", "Heterogénea del Capital",
        "Cálculo Económico", "Teoría del Capital", "Preferencia Temporal", "Productividad Marginal",
        "Interés Natural", "Subsidiaridad", "Humano Acción", "Reconstrucción" 
    ])

    def buscar_informacion(query):
        url = f"https://api.serply.io/v1/scholar/q={query}"
        headers = {
            'X-Api-Key': SERPLY_API_KEY,
            'Content-Type': 'application/json',
            'X-Proxy-Location': 'US',
            'X-User-Agent': 'Mozilla/5.0'
        }
        response = requests.get(url, headers=headers)
        return response.json()

    def generar_definicion(termino, contexto):
        url = "https://api.together.xyz/inference"
        payload = json.dumps({
            "model": "mistralai/Mixtral-8x7B-Instruct-v0.1",
            "prompt": f"Contexto: {contexto}\n\nTérmino: {termino}\n\nProporciona una definición del término económico '{termino}' según la visión de la escuela austríaca de economía. La definición debe ser más larga, detallada, e informativa, similar a una entrada de diccionario extendida. Incluye referencias a fuentes específicas que traten este concepto.\n\nDefinición:",
            "max_tokens": 2048,
            "temperature": 0.7,
            "top_p": 0.7,
            "top_k": 50,
            "repetition_penalty": 1,
            "stop": ["Término:"]
        })
        headers = {
            'Authorization': f'Bearer {TOGETHER_API_KEY}',
            'Content-Type': 'application/json'
        }
        response = requests.post(url, headers=headers, data=payload)
        return response.json()['output']['choices'][0]['text'].strip()

    def create_docx(termino, definicion, fuentes):
        doc = Document()
        doc.add_heading('Diccionario Económico - Escuela Austríaca', 0)

        doc.add_heading('Término', level=1)
        doc.add_paragraph(termino)

        doc.add_heading('Definición', level=2)
        doc.add_paragraph(definicion)

        if fuentes:
            doc.add_heading('Fuentes', level=1)
            for fuente in fuentes:
                doc.add_paragraph(f"{fuente['author']}. ({fuente['year']}). *{fuente['title']}*. {fuente['journal']}, {fuente['volume']}({fuente['issue']}), {fuente['pages']}. {fuente['url']}", style='List Bullet')

        doc.add_paragraph('\nNota: Este documento fue generado por un asistente de IA. Verifica la información con fuentes académicas para un análisis más profundo.')

        return doc

    st.write("Elige un término económico de la lista o propón tu propio término:")

    opcion = st.radio("", ["Elegir de la lista", "Proponer mi propio término"])

    if opcion == "Elegir de la lista":
        termino = st.selectbox("Selecciona un término:", terminos_economicos)
    else:
        termino = st.text_input("Ingresa tu propio término económico:")

    if st.button("Generar entrada de diccionario"):
        if termino:
            with st.spinner("Buscando información y generando definición..."):
                # Buscar información relevante
                resultados_busqueda = buscar_informacion(termino)
                contexto = "\n".join([item["snippet"] for item in resultados_busqueda.get("results", [])])
                fuentes = [{
                    "author": item["author"] if "author" in item else "Autor desconocido",
                    "year": item["year"] if "year" in item else "s.f.",
                    "title": item["title"],
                    "journal": item["journal"] if "journal" in item else "Revista desconocida",
                    "volume": item["volume"] if "volume" in item else "",
                    "issue": item["issue"] if "issue" in item else "",
                    "pages": item["pages"] if "pages" in item else "",
                    "url": item["url"]
                } for item in resultados_busqueda.get("results", [])]

                # Generar definición
                definicion = generar_definicion(termino, contexto)

                # Mostrar la definición
                st.subheader(f"Definición para el término: {termino}")
                st.markdown(f"**{definicion}**")

                # Botón para descargar el documento
                doc = create_docx(termino, definicion, fuentes)
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                st.download_button(
                    label="Descargar definición en DOCX",
                    data=buffer,
                    file_name=f"Definicion_{termino.replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        else:
            st.warning("Por favor, selecciona un término.")
