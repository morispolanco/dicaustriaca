import streamlit as st
import requests
import json
from docx import Document
from io import BytesIO

# Configuración de la página
st.set_page_config(page_title="Diccionario de Términos Socialistas y Marxistas con Refutaciones Filosóficas", page_icon="📚", layout="wide")

# Función para crear la columna de información
def crear_columna_info():
    st.markdown("""
    ## Sobre esta aplicación

    Esta aplicación es un Diccionario de Términos Socialistas y Marxistas con Refutaciones desde una perspectiva filosófica y general. Permite a los usuarios obtener definiciones de términos económicos socialistas o marxistas y sus correspondientes críticas desde un punto de vista amplio y fundamentado.

    ### Cómo usar la aplicación:

    1. Elige un término o tesis socialista/marxista de la lista predefinida o propón tu propio término.
    2. Haz clic en "Obtener definición y refutación" para generar el contenido.
    3. Lee la definición y la refutación filosófica proporcionada.
    4. Si lo deseas, descarga un documento DOCX con toda la información.

    ### Autor y actualización:
    **Moris Polanco**, [Fecha actual]

    ### Cómo citar esta aplicación (formato APA):
    Polanco, M. ([Año actual]). *Diccionario de Términos Socialistas y Marxistas con Refutaciones Filosóficas* [Aplicación web]. [URL de la aplicación]

    ---
    **Nota:** Esta aplicación utiliza inteligencia artificial para generar definiciones y refutaciones basadas en información disponible en línea. Siempre verifica la información con fuentes académicas para un análisis más profundo.
    """)

# Título de la aplicación
st.title("Diccionario de Términos Socialistas y Marxistas con Refutaciones Filosóficas")

# Crear un diseño de dos columnas
col1, col2 = st.columns([1, 2])

# Columna de información
with col1:
    crear_columna_info()

# Columna principal
with col2:
    # Acceder a las claves de API de los secretos de Streamlit
    TOGETHER_API_KEY = st.secrets["TOGETHER_API_KEY"]
    SERPER_API_KEY = st.secrets["SERPER_API_KEY"]

    # Lista de términos y tesis socialistas/marxistas
    terminos_socialistas = [
        "Lucha de clases", "Plusvalía", "Alienación", "Materialismo histórico", "Dictadura del proletariado",
        "Modo de producción", "Socialismo científico", "Revolución proletaria", "Conciencia de clase",
        "Imperialismo", "Capital constante y variable", "Fetichismo de la mercancía", "Acumulación primitiva",
        "Ejército industrial de reserva", "Superestructura e infraestructura", "Socialización de los medios de producción",
        "Teoría del valor-trabajo", "Contradicciones del capitalismo", "Comunismo primitivo",
        "Internacionalismo proletario", "Determinismo económico", "Dialéctica materialista",
        "Explotación laboral", "Pauperización", "Concentración del capital"
    ]

    def buscar_informacion(query):
        url = "https://google.serper.dev/search"
        payload = json.dumps({
            "q": f"{query} socialismo marxismo"
        })
        headers = {
            'X-API-KEY': SERPER_API_KEY,
            'Content-Type': 'application/json'
        }
        response = requests.post(url, headers=headers, data=payload)
        if response.status_code != 200:
            st.error(f"Error en la búsqueda: {response.status_code} - {response.text}")
            return None
        return response.json()

    def generar_definicion_y_refutacion(termino, contexto):
        url = "https://api.together.xyz/inference"
        payload = json.dumps({
            "model": "mistralai/Mixtral-8x7B-Instruct-v0.1",
            "prompt": f"""Contexto: {contexto}

Término: {termino}

1. Proporciona una definición concisa pero informativa del término o tesis socialista/marxista '{termino}', similar a una entrada de diccionario.

2. Luego, proporciona una refutación o crítica amplia y fundamentada desde un punto de vista general o filosófico. Esta refutación debe:
   - Abordar los principios fundamentales y las implicaciones filosóficas del concepto.
   - Considerar aspectos éticos, políticos y sociales más allá de lo puramente económico.
   - Presentar argumentos lógicos y ejemplos históricos o teóricos relevantes.
   - Explorar las posibles contradicciones o debilidades en la teoría.
   - Ofrecer perspectivas alternativas de diferentes escuelas de pensamiento.

La refutación debe ser equilibrada, académica y basada en un análisis crítico profundo.

Definición:

Refutación filosófica:""",
            "max_tokens": 2048,
            "temperature": 0,
            "top_p": 0.7,
            "top_k": 50,
            "repetition_penalty": 0,
            "stop": ["Término:"]
        })
        headers = {
            'Authorization': f'Bearer {TOGETHER_API_KEY}',
            'Content-Type': 'application/json'
        }
        response = requests.post(url, headers=headers, data=payload)

        if response.status_code != 200:
            st.error(f"Error en la API de Together: {response.status_code} - {response.text}")
            return None

        try:
            return response.json()['output']['choices'][0]['text'].strip()
        except KeyError as e:
            st.error(f"Error al procesar la respuesta de la API de Together: {e}")
            st.json(response.json())  # Muestra la respuesta completa para debug
            return None

    def create_docx(termino, definicion, refutacion, fuentes):
        doc = Document()
        doc.add_heading('Diccionario de Términos Socialistas y Marxistas con Refutaciones Filosóficas', 0)

        doc.add_heading('Término', level=1)
        doc.add_paragraph(termino)

        doc.add_heading('Definición', level=2)
        doc.add_paragraph(definicion)

        doc.add_heading('Refutación Filosófica', level=2)
        doc.add_paragraph(refutacion)

        doc.add_heading('Fuentes', level=1)
        for fuente in fuentes:
            doc.add_paragraph(fuente, style='List Bullet')

        doc.add_paragraph('\nNota: Este documento fue generado por un asistente de IA. Verifica la información con fuentes académicas para un análisis más profundo.')

        return doc

    # Interfaz de usuario
    st.write("Elige un término o tesis socialista/marxista de la lista o propón tu propio término:")

    opcion = st.radio("", ["Elegir de la lista", "Proponer mi propio término"])

    if opcion == "Elegir de la lista":
        termino = st.selectbox("Selecciona un término:", terminos_socialistas)
    else:
        termino = st.text_input("Ingresa tu propio término o tesis socialista/marxista:")

    if st.button("Obtener definición y refutación"):
        if termino:
            with st.spinner("Buscando información y generando contenido..."):
                # Buscar información relevante
                resultados_busqueda = buscar_informacion(termino)
                if resultados_busqueda:
                    contexto = "\n".join([item["snippet"] for item in resultados_busqueda.get("organic", []) if "snippet" in item])
                    fuentes = [item["link"] for item in resultados_busqueda.get("organic", [])]

                    # Generar definición y refutación
                    contenido = generar_definicion_y_refutacion(termino, contexto)

                    if contenido:
                        # Dividir el contenido en definición y refutación
                        partes = contenido.split("Refutación filosófica:")
                        if len(partes) == 2:
                            definicion, refutacion = partes
                        else:
                            st.error("No se pudo separar la definición de la refutación.")
                            st.text(contenido)  # Mostrar el contenido completo para debug
                            definicion = contenido
                            refutacion = "No se pudo generar una refutación."

                        # Mostrar el contenido
                        st.subheader(f"Término: {termino}")
                        st.markdown("**Definición:**")
                        st.write(definicion.strip())
                        st.markdown("**Refutación filosófica:**")
                        st.write(refutacion.strip())

                        # Botón para descargar el documento
                        doc = create_docx(termino, definicion.strip(), refutacion.strip(), fuentes)
                        buffer = BytesIO()
                        doc.save(buffer)
                        buffer.seek(0)
                        st.download_button(
                            label="Descargar contenido en DOCX",
                            data=buffer,
                            file_name=f"Definicion_y_Refutacion_{termino.replace(' ', '_')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    else:
                        st.error("No se pudo generar el contenido. Por favor, intenta de nuevo.")
                else:
                    st.error("No se pudo obtener información relevante. Por favor, intenta de nuevo.")
        else:
            st.warning("Por favor, selecciona o ingresa un término.")
