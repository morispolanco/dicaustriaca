import streamlit as st
import requests
import json
from docx import Document
from io import BytesIO

# Set page configuration
st.set_page_config(page_title="Diccionario de Economía Austríaca", page_icon="📊", layout="wide")

# Function to create the information column
def crear_columna_info():
    st.markdown("""
    ## Sobre esta aplicación

    Esta aplicación es un Diccionario de Economía basado en la visión de la Escuela Austríaca. Permite a los usuarios obtener definiciones de términos económicos según la interpretación de esta escuela de pensamiento.

    ### Cómo usar la aplicación:

    1. Elija un término económico de la lista predefinida.
    2. Haga clic en "Generar entrada de diccionario" para obtener la definición desde la perspectiva de la Escuela Austríaca.
    3. Lea la definición y las fuentes proporcionadas.
    4. Si lo desea, descargue un documento DOCX con toda la información.

    ### Autor y actualización:
    **Moris Polanco**, 28 ag 2024

    ### Cómo citar esta aplicación (formato APA):
    Polanco, M. (2024). *Diccionario de Economía Austríaca* [Aplicación web]. https://economiaaustriaca.streamlit.app

    ---
    **Nota:** Esta aplicación utiliza inteligencia artificial para generar definiciones basadas en la visión de la Escuela Austríaca. Verifique la información con fuentes adicionales para un análisis más profundo.
    """)

# Titles and Main Column
st.title("Diccionario de Economía Austríaca")

col1, col2 = st.columns([1, 2])

with col1:
    crear_columna_info()

with col2:
    TOGETHER_API_KEY = st.secrets["TOGETHER_API_KEY"]
    SERPLY_API_KEY = st.secrets["SERPLY_API_KEY"]

    # 101 economic terms related to the Austrian School perspective
    terminos_economicos = sorted([
        "Acción humana", "Ahorro", "Anarcocapitalismo", "Arbitraje", "Banco central", "Banca de reserva fraccionaria",
        "Bienes de capital", "Bienes de consumo", "Cálculo económico", "Capitalismo", "Ciclo económico austriaco",
        "Competencia", "Consumo", "Coste de oportunidad", "Crédito", "Deflación", "Demanda", "Depresión económica",
        "Derechos de propiedad", "Dinero fiduciario", "Dinero mercancía", "División del trabajo", "Economía de mercado",
        "Efecto Ricardo", "Empresario", "Escasez", "Escuela de Salamanca", "Estado", "Estructura del capital",
        "Eviccionismo", "Externalidad", "Frontera de posibilidades de producción", "Gobierno", "Hiperinflación",
        "Homo agens", "Imperialismo monetario", "Incentivos", "Incertidumbre", "Indexación", "Inflación",
        "Intervención estatal", "Inversión", "Laissez-faire", "Ley de asociación de Ricardo", "Ley de costos",
        "Ley de la utilidad marginal decreciente", "Ley de oferta y demanda", "Ley de preferencia temporal",
        "Ley de rendimientos decrecientes", "Ley de Say", "Liberalismo clásico", "Libre mercado", "Liquidez",
        "Mano invisible", "Marginalismo", "Medios de producción", "Mercado negro", "Metodología apriorística",
        "Monopolio", "Orden espontáneo", "Patrón oro", "Planificación central", "Poder adquisitivo", "Precios",
        "Preferencia temporal", "Praxeología", "Privatización", "Producción", "Proteccionismo", "Punto de equilibrio",
        "Racionamiento", "Recesión", "Riesgo moral", "Salarios", "Satisfacción de necesidades", "Sector privado",
        "Sector público", "Selección natural económica", "Socialismo", "Soberanía del consumidor", "Subjetivismo",
        "Tasa de interés natural", "Teorema de la imposibilidad del socialismo", "Teorema de regresión",
        "Teoría austriaca del ciclo económico", "Teoría del valor subjetivo", "Tiempo", "Tipos de interés",
        "Utilidad marginal", "Valor", "Valor presente", "Velocidad de circulación del dinero", "Ventaja comparativa",
        "Voluntarismo"
    ])

    def buscar_informacion(query):
        url = f"https://api.serply.io/v1/scholar/q={query} Austrian School of Economics"
        headers = {
            'X-Api-Key': SERPLY_API_KEY,
            'Content-Type': 'application/json',
            'X-Proxy-Location': 'US',
            'X-User-Agent': 'Mozilla/5.0'
        }
        response = requests.get(url, headers=headers)
        return response.json()

    def generar_definicion(termino, contexto):
        url = "https://api.together.xyz/inference"
        payload = json.dumps({
            "model": "mistralai/Mixtral-8x7B-Instruct-v0.1",
            "prompt": f"Contexto: {contexto}\n\nTérmino: {termino}\n\nProporciona una definición del término económico '{termino}' según la visión de la Escuela Austríaca de Economía. La definición debe ser detallada e informativa, similar a una entrada de diccionario extendida. Incluye referencias a economistas austriacos relevantes y conceptos relacionados.\n\nDefinición:",
            "max_tokens": 2048,
            "temperature": 0.7,
            "top_p": 0.7,
            "top_k": 50,
            "repetition_penalty": 1,
            "stop": ["Término:"]
        })
        headers = {
            'Authorization': f'Bearer {TOGETHER_API_KEY}',
            'Content-Type': 'application/json'
        }
        response = requests.post(url, headers=headers, data=payload)
        return response.json()['output']['choices'][0]['text'].strip()

    def create_docx(termino, definicion, fuentes):
        doc = Document()
        doc.add_heading('Diccionario de Economía Austríaca', 0)

        doc.add_heading('Término', level=1)
        doc.add_paragraph(termino)

        doc.add_heading('Definición', level=2)
        doc.add_paragraph(definicion)

        # Agregar "Fuentes" solo si hay fuentes disponibles
        if fuentes:
            doc.add_heading('Fuentes', level=1)
            for fuente in fuentes:
                doc.add_paragraph(f"{fuente['author']}. ({fuente['year']}). *{fuente['title']}*. {fuente['journal']}, {fuente['volume']}({fuente['issue']}), {fuente['pages']}. {fuente['url']}", style='List Bullet')

        doc.add_paragraph('\nNota: Este documento fue generado por un asistente de IA. Verifica la información con fuentes académicas para un análisis más profundo.')

        return doc

    st.write("Elige un término económico de la lista:")

    termino = st.selectbox("Selecciona un término:", terminos_economicos)

    if st.button("Generar entrada de diccionario"):
        if termino:
            with st.spinner("Buscando información y generando definición..."):
                # Buscar información relevante
                resultados_busqueda = buscar_informacion(termino)
                contexto = "\n".join([item["snippet"] for item in resultados_busqueda.get("results", [])])
                fuentes = [{
                    "author": item["author"] if "author" in item else "Autor desconocido",
                    "year": item["year"] if "year" in item else "s.f.",
                    "title": item["title"],
                    "journal": item["journal"] if "journal" in item else "Revista desconocida",
                    "volume": item["volume"] if "volume" in item else "",
                    "issue": item["issue"] if "issue" in item else "",
                    "pages": item["pages"] if "pages" in item else "",
                    "url": item["url"]
                } for item in resultados_busqueda.get("results", [])]

                # Generar definición
                definicion = generar_definicion(termino, contexto)

                # Mostrar la definición
                st.subheader(f"Definición para el término: {termino}")
                st.markdown(f"**{definicion}**")

                # Botón para descargar el documento
                doc = create_docx(termino, definicion, fuentes)
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                st.download_button(
                    label="Descargar definición en DOCX",
                    data=buffer,
                    file_name=f"Definicion_{termino.replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        else:
            st.warning("Por favor, selecciona un término.")
